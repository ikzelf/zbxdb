#!/usr/bin/env python
""" read given dashboard, get the widgets, get the graphs and
    dump them in a .docx document
    specify report period with -S -1,08:00:00 -P 36000
                   for yesterday from 08:00:00 for 36000 seconds
                   or simply   -S yyyy-mm-dd 08:00:00 -P 36000
"""
import io
import json
import logging
import optparse
import sys
import time
from calendar import monthrange
from datetime import datetime, timedelta

import requests
from dateutil.relativedelta import relativedelta
from docx import Document
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.shared import Cm
from PIL import Image
from pyzabbix import ZabbixAPI, ZabbixAPIException


def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink

def api_connect():
    zapi = ZabbixAPI(options.protocol + "://" + options.server + options.path)
    zapi.login(options.username, options.password)

    return zapi

def dash_lookup(name):
    return zapi.dashboard.get(search={"name": name}, output="extend")

def host_lookup(graphid):

    gitem = zapi.graphitem.get( graphids=graphid
                               ,output="extend"
                              )
    log.debug(gitem)
    itemid = gitem[0]['itemid']
    log.debug("itemid: {}".format(itemid))
    item = zapi.item.get( itemids=itemid
                         ,output="extend"
                         )
    log.debug(item)
    hostid = item[0]['hostid']

    return zapi.host.get(hostids=hostid,
                           output="extend",
                           selectInventory=['name','site_address_b','asset_tag'],
                           )

def interfaces_lookup(item_id):
    ifacelist = zapi.hostinterface.get(output='extend', filter={"hostid": item_id,"main": 1} )

    return ifacelist

def widgets_lookup(id):
    return zapi.dashboard.get(dashboardids=id,
                               output="extend",
                               selectWidgets="extend"
                               )
def get_graphs(ids):
    log.debug("getting graphs {}".format(ids))

    return zapi.graph.get(output="extend",
                          graphids=ids)


parser = optparse.OptionParser()
parser.add_option('-u', '--username', help="Username", default="admin")
parser.add_option('-p', '--password', help="Password", default="zabbix")
parser.add_option('-r', '--protocol', help="Protocol to be used", default="https")
parser.add_option('-s', '--server', help="Host To talk to the web api", default="your_zabbix_server")
parser.add_option('-d', '--path', help="Path", default="")
parser.add_option('-f', '--file', help="docx file for input template", default="template.docx")
parser.add_option('-D', '--dashboard', help="dashboard to be dumped")
parser.add_option('-S', '--starttime', help="starttime", default="-1d,08:00:00")
parser.add_option('-P', '--period', help="period xxM or xxd or xxh or xxm or xxs", default="10h")
parser.add_option('-V', '--verify', help="check ssl certificate", default=True)
parser.add_option('-W', '--width', help="image width", default="")
parser.add_option('-H', '--height', help="image height", default="")
parser.add_option('-v', '--verbosity', help="control logging", action="count", default=0)
(options,args)=parser.parse_args()
stream = logging.StreamHandler(sys.stdout)
stream.setLevel(logging.DEBUG)
log = logging.getLogger('pyzabbix')
log.addHandler(stream)
api = options.protocol + "://" + options.server + options.path

if options.verbosity > 0:
    log.setLevel(logging.DEBUG)

log.debug(options)

if len(options.starttime.split(",")) == 2:
    reldate, reptime = options.starttime.split(",")
    relunit = reldate[-1]

    if relunit == 'd':
        # starting "hh24:mi:ss" is specified
        starttime=datetime.now() + timedelta(days=int(reldate[:-1]))
        starttimes="{} {}".format(starttime.strftime("%Y-%m-%d"), reptime)
    elif relunit == 'M':
        # starting day in month is specified
        starttime = datetime.now().replace(day=int(reptime))+ relativedelta(months=int(reldate[:-1]))
        starttimes="{} {}".format(starttime.strftime("%Y-%m-%d"), "00:00:00")
        log.debug("starttime: {}".format(starttimes))
else:
    starttimes=options.starttime
dt = datetime.strptime(starttimes, '%Y-%m-%d %H:%M:%S')
stime= time.mktime(dt.timetuple())

duration=int(options.period[:-1])
durationUnit = options.period[-1]
log.debug("duration: {}, unit: {}".format(duration, durationUnit))

if durationUnit == "M":
    log.debug(type(stime))
    log.debug("stime: {}".format(stime))
    etime = starttime + relativedelta(months=duration)
    log.debug(type(etime))
    log.debug("etime: {}".format(etime))
    log.debug(etime.timestamp())
    period = etime.timestamp() - stime
elif durationUnit == "d":
    period = duration * 24 * 3600
elif durationUnit == "h":
    period = duration * 3600
elif durationUnit == "m":
    period = duration * 60
else:
    period = duration
etime = stime + int(period)

log.warning("generating for {} ({}) period {} to {}".format(
    datetime.fromtimestamp(stime).strftime('%Y-%m-%d %H:%M:%S'),
    stime,
    options.period,
    datetime.fromtimestamp(etime).strftime('%Y-%m-%d %H:%M:%S'),
    ))

if not options.dashboard:
    print("specify dashboard name to dump",file=sys.stderr)
    sys.exit(1)

if not options.file:
    print("specify file to write docx to",file=sys.stderr)
    sys.exit(1)
print("Using {} to dump dashboard \"{}\"\nfrom url {}://{}{}".format(options.username,
    options.dashboard,options.protocol, options.server, options.path))


zapi = api_connect()

dash =  dash_lookup(options.dashboard)
log.debug("dashboard[s] found: {}".format(dash))
dashid = dash[0]['dashboardid']
log.debug("using first one found: {}".format(dashid))
widgets= widgets_lookup(dashid)
log.debug("widgets found: {}".format(widgets))
graphids=[]

for widget in widgets[0]['widgets']:
    # order is kind of random

    if widget['type'] in ["graph"]:
        log.debug("using widget : {} type {} x {} y {} name {}".format(
            widget['widgetid'],widget['type'], widget['x'],
            widget['y'],
            widget['name']))

        if widget['fields'][0]['name'] == "graphid":
            graphids.append({"graphid": widget['fields'][0]['value'],
                                "name": widget['name'],
                                "y": widget['y'],
                                "x": widget['x']
                            })
        else:
            log.warning("ignoring widget {} field {}".format(
                widget['widgetid'], widget['fields']))
    else:
        log.warning("ignoring widget : {} type {} not [svg]graph field {}".format(
            widget['widgetid'],widget['type'], widget['fields'][0]))

sgraphids = sorted(graphids, key=lambda k: (int(k['x']), int(k['y'])) )

for g in sgraphids:
    log.debug(g)
# Set login URL for the Frontend (frontend access is needed, as we
# cannot retrieve graph images via the API)
loginurl = api + "/index.php"
# Data that needs to be posted to the Frontend to log in
logindata = {'autologin' : '1', 'name' : options.username, 'password' :
        options.password, 'enter' : 'Sign in'}
# We need to fool the frontend into thinking we are a real browser
headers = {'User-Agent' : 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_4) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/81.0.4044.129 Safari/537.36', 'Content-type' : 'application/x-www-form-urlencoded'}

# setup a session object so we can reuse session cookies
session=requests.session()
# Login to the frontend
login=session.post(loginurl, params=logindata, headers=headers, verify=options.verify)
document = Document(options.file)
paragraph = document.add_paragraph(
        """Rapport gegenereerd op {}
Rapportage periode van {} tot {}""".format(
        datetime.now(),
        datetime.fromtimestamp(stime).strftime('%Y-%m-%d %H:%M:%S'),
        datetime.fromtimestamp(etime).strftime('%Y-%m-%d %H:%M:%S'),
        ))

document.add_page_break()
document.add_heading('De Grafieken', level=1)

if session.cookies['zbx_sessionid']:
    for g in sgraphids:
        log.debug(g)
        graph = get_graphs(g['graphid'])[0]
        log.debug(graph)
        # Select the right graph generator according to graph type
        # type 3 = Exploded graph

        if graph['graphtype'] == "3":
            generator = "chart6.php"
        # type 2 = Pie graph
        elif graph['graphtype'] == "2":
            generator = "chart6.php"
        # type 1 = Stacked graph
        elif graph['graphtype'] == "1":
            generator = "chart2.php"
        # type 0 = Normal graph
        elif graph['graphtype'] == "0":
            generator = "chart2.php"
        # catch-all in case someone invents a new type/generator
        else:
            generator = "chart2.php"
        log.debug("generator {}".format(generator))
        # Set width and height

        if options.width:
            width = options.width
        else:
            width = graph['width']

        if options.height:
            height = options.height
        else:
            height = graph['height']
        # Build the request for the graph
        # v4.4 chart2.php/graphid=5584&from=2020-04-29%2008:00:00&to=2020-04-29%2018:00:00&profileIdx=web.graphs.filter&profileIdx2=2206&width=600&height=201&_=u84i64mh&screenid=
        graphurl = api + "/" + generator + "?graphid=" + str(graph['graphid']) \
        + "&from=" + datetime.fromtimestamp(stime).strftime('%Y-%m-%d %H:%M:%S') \
        + "&to=" + datetime.fromtimestamp(etime).strftime('%Y-%m-%d %H:%M:%S') \
        + "&width=" + str(width) \
        + "&height=" + str(height) \
        + "&profileIdx=web.graphs.filter"
        log.warning("x {} y {} url {}".format(g['x'],g['y'],graphurl))
        # get the graph
        graphreq = session.get(graphurl,verify=options.verify)

        hosts = (host_lookup(graph['graphid']))

        for host in hosts:
            log.debug(host['inventory'])

        # read the data as an image
        graphpng = io.BytesIO(graphreq.content)

        paragraph_h = document.add_heading("device: {} locatie: {}".format(
                host['inventory']['name'],
                host['inventory']['site_address_b']
                )
                , level=2)
        # paragraph_h = document.add_paragraph("device: {} locatie: {}".format(
                # host['inventory']['name'],
                # host['inventory']['site_address_b']
                # ))
        paragraph_h.paragraph_format.keep_with_next = True
        x = document.add_picture(graphpng, width=Cm(18), height=None)
        paragraph_t = document.add_paragraph('')

docname=options.dashboard +"-"+ \
    datetime.fromtimestamp(stime).strftime('%Y-%m-%d %H:%M:%S-')+str(options.period)+".docx"
f = open(docname, "w")
f.close()
document.save(docname)
